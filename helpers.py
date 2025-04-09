from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import json
import uuid
import random, re
import pandas as pd
from chromadb import PersistentClient
from chromadb.config import Settings
from chromadb.utils import embedding_functions
from classes import SubmittalPage, BOQItem, ItemVariant
from pydantic import ValidationError
from typing import List

# --- DOCS IMPORTS --
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import markdown
from bs4 import BeautifulSoup


def read_pdf(file_path):
    docs_reader = PyPDFLoader(file_path)
    docs = docs_reader.load()

    return docs

def extract_data_from_doc(file_path):
    pages = []
    docs = read_pdf(file_path)
    max_length = get_docs_max_length(docs)
    for page in docs:
      page_num = page.metadata['page'] + 1
      #Chunking
      page_chunks = text_chunks(page.page_content,max_length)


      #Creating Page Object
      contract = SubmittalPage(
          source = page.metadata['source'],
          page_num = page_num,
          content = page_chunks,
      )
      pages.append(contract)

    return pages

def get_docs_max_length(docs):
    #Check for max_length of single page
    max_length = 0
    for i in range(len(docs)):
        if len(docs[i].page_content) > max_length:
            max_length = len(docs[i].page_content)
            return max_length
        else:
            continue

def text_chunks(text,max_length):
  chunks = {}
  #Configure Langchain text splitter
  splitter_config = {
      "chunk_size": round((max_length/3),0),
      "chunk_overlap": max_length/10
  }
  text_splitter = RecursiveCharacterTextSplitter(**splitter_config)
  chunks=text_splitter.split_text(text)
  return chunks


def save_to_json(extracted_data):
    with open("processed_docs/processed_data.json","w",encoding= 'utf-8') as f:
        result = json.dumps(extracted_data,indent=4,ensure_ascii=False)
        f.write(result)

def save_pdf_to_collection(extracted_data,persistance_path,collection_name):
    client = PersistentClient(persistance_path)
    embedding_function = embedding_functions.SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")
    collection = client.get_or_create_collection(name=collection_name,embedding_function=embedding_function)

    for id,page in enumerate(extracted_data):
        if page.content == []:
            continue
        else:
            pages = page.content
            ids = [str(uuid.uuid4()) for _ in pages]
            source = extracted_data[id].source.split("/")[-1]
            metadatas = [{"source":source, "page": page_num } for page_num in range(1, len(pages) + 1)]
            collection.upsert(
                documents=pages,
                ids=ids,
                metadatas=metadatas
            )

def save_boq_to_collection(extracted_data,persistance_path,collection_name):
    client = PersistentClient(persistance_path)
    embedding_function = embedding_functions.SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")
    collection = client.get_or_create_collection(name=collection_name,embedding_function=embedding_function)

    metadatas = [{"item_num":item.itemNum, "variants_num": len(item.variants)} for item in extracted_data]
    ids = [str(uuid.uuid4()) for _ in extracted_data]
    documents = [item.model_dump_json() for item in extracted_data]

    collection.upsert(
        documents=documents,
        ids=ids,
        metadatas=metadatas
    )

#Processing Excel files (BOQs)
def read_boq(file_path: str, sheet_name: str = None) -> pd.DataFrame:
    """Reads an Excel file and returns a pandas DataFrame."""
    try:
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        if df.empty:
            raise pd.errors.EmptyDataError(f"The file or sheet '{sheet_name}' is empty.")
        return df
    except FileNotFoundError:
        raise FileNotFoundError(f"The file '{file_path}' was not found.")
    except pd.errors.EmptyDataError:
        raise pd.errors.EmptyDataError(f"The file or sheet '{sheet_name}' is empty.")
    except Exception as e:
        raise Exception(f"An error occurred while reading the file: {e}")


def process_data(df: pd.DataFrame) -> pd.DataFrame:
    """Cleans and transforms the DataFrame."""
    try:
        df = df.fillna("")
        if len(df.columns) < 6:
           raise ValueError("Input DataFrame does not have enough columns. Expected at least 6.")

        df.columns = ["Item", "Description", "Unit", "QTY", "Unit Price", "Total price"]
        #df = df.iloc[1:].reset_index(drop=True)
        df = df[df["Description"] != ""].reset_index(drop=True)
        df["Unit Price"] = [random.randint(500, 1000) for i in range(len(df["Unit Price"]))]
        df["Unit Price"] = pd.to_numeric(df["Unit Price"], errors='coerce').fillna(0.0)
        df["QTY"] = pd.to_numeric(df["QTY"], errors='coerce').fillna(0.0)
        df["Total price"] = df["QTY"] * df["Unit Price"]

        return df

    except KeyError as e:
        raise ValueError(f"Missing expected column(s) in DataFrame: {e}")
    except Exception as e:
        raise Exception(f"An error occurred during data processing: {e}")

def filter_chromadb_query_results(results,degree):
    filtered_results = []
    for i in range(len(results['distances'])):
        for j in range(len(results['distances'][i])):
            if results['distances'][i][j] < degree:
                filtered_results.append({
                    "id": results['ids'][i][j],
                    "document": results['documents'][i][j],
                    "metadata": results['metadatas'][i][j],
                    "distance": results['distances'][i][j]
                })
    return filtered_results


def retrieve_from_vectorstore(persistent_path:str, boq_collection_name: str, specs_collection_name: str,summarized_text: str):
  settings = Settings(anonymized_telemetry=False)
  client = PersistentClient(persistent_path,settings=settings)
  boq_collection = client.get_collection(name=boq_collection_name)
  specs_collection = client.get_collection(name=specs_collection_name)
  query_text = summarized_text

  results_specs = specs_collection.query(
      query_texts=[query_text],
      n_results=5
  )

  results_boq = boq_collection.query(
      query_texts=[query_text],
      n_results=4,

  )

  query_boq_text = [results_boq['documents'][0][i] for i in range(len(results_boq['documents'][0]))]
  query_specs_text = [results_specs['documents'][0][i] for i in range(len(results_specs['documents'][0]))]
  return query_boq_text, query_specs_text

def datasheet_content(data_sheet_file_path, model):
  submittal_data_text = PyPDFLoader(data_sheet_file_path).load()
  submittal_data_text = [submittal_data_text[i].page_content for i in range(len(submittal_data_text))]
  submittal_data_text = "\n".join(submittal_data_text)

  data_sheet_summarize_query = """
  You are a helpful assistant, I want you to summarize the following text and return only the summarized text without any additional texts in the beginning or at the end.

  Text : {text}

  Note that you should extract the following information from the text:
  1-specifications
  2-design characteristics
  3-general features
  4-data sheets
  5-any important or valuable info

  **Do not miss any information ( the summarization purpose is to organize the information in the datasheet for further processing
  """.format(text = "\n".join(submittal_data_text))

  summarized_text = model.invoke(data_sheet_summarize_query).content
  return submittal_data_text

def group_boq_items(df: pd.DataFrame, delimiter: str = "-") -> pd.core.groupby.DataFrameGroupBy:
    """Groups DataFrame rows by item number, using regex for more robust matching."""
    try:
        df["item_group_num"] = ""
        # Regex to match the first part of the item number (digits before the delimiter)
        item_number_pattern = re.compile(r"^(\d+)")

        for index, row in df.iterrows():
            item_str = str(row["Item"])
            if item_str != "":
                match = item_number_pattern.match(item_str)
                if match:
                    number_str = match.group(1)  # Get the captured group (the digits)
                    try:
                        df.loc[index, "item_group_num"] = int(number_str)
                    except ValueError:
                        df.loc[index, "item_group_num"] = -1
                        print(f"Warning: Invalid item number format at row {index}: {number_str}")
                else:
                    df.loc[index, "item_group_num"] = -1
                    print(f"Warning: Invalid item number format at row {index}: {item_str}")

        df_group = df.groupby("item_group_num")
        return df_group

    except KeyError as e:
        raise ValueError(f"Missing expected column(s) during grouping: {e}")
    except Exception as e:
        raise Exception(f"An error occurred during item grouping: {e}")


def extract_boq_data(df_group: pd.core.groupby.DataFrameGroupBy) -> List[BOQItem]:
    """Extracts BOQ data from grouped DataFrame."""
    boq_item_list: List[BOQItem] = []
    try:
        for item_group_num, item_group_df in df_group:
            variant_list: List[ItemVariant] = []
            if len(item_group_df) > 1:
                for i in range(1, len(item_group_df)):
                    row = item_group_df.iloc[i]
                    try:
                        item_variant = ItemVariant(
                            variant_num=str(row["Item"]),
                            description=str(row["Description"]),
                            unit=str(row["Unit"]),
                            quantity=row["QTY"],
                            rate=row["Unit Price"],
                            amount=row["Total price"]
                        )
                        variant_list.append(item_variant)
                    except ValidationError as ve:
                        print(
                            f"Validation error for ItemVariant at row {i} in group {item_group_num}: {ve}")
                        continue
                    except Exception as e:
                        print(f"Error creating ItemVariant at row {i} in group {item_group_num}: {e}")
                        continue

                main_row = item_group_df.iloc[0]
                try:
                    boq_item = BOQItem(
                        itemNum=str(main_row["Item"]),
                        description=list(item_group_df["Description"].astype(str)),
                        unit=str(main_row["Unit"]),
                        quantity=main_row["QTY"],
                        rate=main_row["Unit Price"],
                        amount=main_row["Total price"],
                        variants=variant_list
                    )
                    boq_item_list.append(boq_item)
                except ValidationError as ve:
                    print(f"Validation error for BOQItem in group {item_group_num}: {ve}")
                except Exception as e:
                    print(f"Error creating BOQItem in group {item_group_num}: {e}")


            else:
                main_row = item_group_df.iloc[0]
                try:
                    boq_item = BOQItem(
                        itemNum=str(main_row["Item"]),
                        description=list(item_group_df["Description"].astype(str)),
                        unit=str(main_row["Unit"]),
                        quantity=main_row["QTY"],
                        rate=main_row["Unit Price"],
                        amount=main_row["Total price"],
                        variants=[]
                    )
                    boq_item_list.append(boq_item)
                except ValidationError as ve:
                    print(f"Validation error for BOQItem in single-line group {item_group_num}: {ve}")
                except Exception as e:
                    print(f"Error creating BOQItem in single-line group {item_group_num}: {e}")

    except Exception as e:
        raise Exception(f"An unexpected error occurred during data extraction: {e}")

    return boq_item_list

def generate_report(boq_retrieved_docs,specs_retrieved_docs, summarized_text, model):
  with open("prompt_gallery\generate_report.txt","r") as f:
    query = f.read()
  query = query.format(specs = specs_retrieved_docs, boq = boq_retrieved_docs, submittal_text = summarized_text)
  model_response = model.invoke(query)
  return model_response

def create_docx_from_markdown(markdown_text, output_filename):
    """
    Creates a .docx file from the provided markdown text with landscape layout,
    modern styles, margins, adjusted line spacing, and a page border.

    Args:
        markdown_text: The markdown text to convert.
        output_filename: The name of the output .docx file or a file-like object.
        
    Returns:
        Path to the created document or the file-like object.
    """
    # Convert markdown to HTML for better parsing
    html = markdown.markdown(markdown_text, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'html.parser')
    
    document = Document()

    # --- Document Setup (Landscape, Margins, Styles) ---
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)  # A4 Landscape width
    section.page_height = Inches(8.27)  # A4 Landscape height
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Configure styles
    _configure_document_styles(document)
    
    # Add page border
    _add_page_border(document)
    
    # Parse HTML and create document content
    _process_html_elements(soup, document)
    
    # Save the document
    document.save(output_filename)
    return output_filename


def _configure_document_styles(document):
    """Configure document styles for modern appearance."""
    # Normal text style
    normal_style = document.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_paragraph_format = normal_style.paragraph_format
    normal_paragraph_format.space_before = Pt(0)
    normal_paragraph_format.space_after = Pt(0)
    normal_paragraph_format.line_spacing = 1.15
    
    # Heading styles
    for level in range(1, 6):
        heading_style_name = f'Heading {level}'
        if heading_style_name in document.styles:
            heading_style = document.styles[heading_style_name]
            heading_style.font.name = 'Calibri'
            heading_style.font.bold = True
            heading_style.font.size = Pt(14 - level)  # Decrease size as level increases
            heading_style.paragraph_format.space_before = Pt(12 - level)
            heading_style.paragraph_format.space_after = Pt(6)
    
    # Table style
    table_style = document.styles.add_style('CustomTableStyle', WD_STYLE_TYPE.TABLE)
    table_style.base_style = document.styles['Table Grid']
    
    # Code style for code blocks
    code_style = document.styles.add_style('CodeStyle', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(10)
    code_paragraph_format = code_style.paragraph_format
    code_paragraph_format.space_before = Pt(6)
    code_paragraph_format.space_after = Pt(6)
    
    # List style
    list_style = document.styles.add_style('ListStyle', WD_STYLE_TYPE.PARAGRAPH)
    list_style.base_style = document.styles['Normal']
    list_paragraph_format = list_style.paragraph_format
    list_paragraph_format.left_indent = Inches(0.25)
    list_paragraph_format.first_line_indent = Inches(-0.25)


def _add_page_border(document):
    """Add a page border to the document."""
    # Create the border element and attributes
    sections = document.sections
    for section in sections:
        sectPr = section._sectPr
        
        # Create page borders element
        pg_borders = OxmlElement('w:pgBorders')
        pg_borders.set(qn('w:offsetFrom'), 'page')
        
        # Create the border elements
        border_names = ['top', 'left', 'bottom', 'right']
        for border_name in border_names:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), 'auto')
            pg_borders.append(border)
        
        # Add borders to section properties
        sectPr.append(pg_borders)


def _process_html_elements(soup, document):
    """Process HTML elements and convert them to Word document elements."""
    # Process the body content
    if soup.body:
        for element in soup.body.contents:
            _process_element(element, document)
    else:
        # If no body found, process all top-level elements
        for element in soup.contents:
            _process_element(element, document)


def _process_element(element, document):
    """Process a single HTML element and its children."""
    # Skip if element is just a newline or whitespace
    if element.name is None:  # It's a NavigableString
        if element.strip():  # Only add non-empty strings
            p = document.add_paragraph(style='Normal')
            p.add_run(element.strip())
        return
    
    tag_name = element.name
    
    if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(tag_name[1])
        heading = document.add_heading(element.get_text(), level=level)
        
    elif tag_name == 'p':
        if element.find('code') and len(list(element.children)) == 1:
            # This is probably a code block
            p = document.add_paragraph(style='CodeStyle')
            p.add_run(element.get_text().strip())
        else:
            p = document.add_paragraph(style='Normal')
            # Process inline formatting
            _process_inline_elements(element, p)
            
    elif tag_name == 'ul' or tag_name == 'ol':
        _process_list(element, document)
        
    elif tag_name == 'table':
        _process_table(element, document)
        
    elif tag_name == 'pre':
        # Code block
        code_text = element.get_text().strip()
        p = document.add_paragraph(style='CodeStyle')
        p.add_run(code_text)
        
    elif tag_name == 'hr':
        document.add_paragraph('_' * 50, style='Normal')
        
    elif tag_name == 'blockquote':
        for child in element.contents:
            if child.name:
                _process_element(child, document)
            elif child.string and child.string.strip():
                p = document.add_paragraph(style='Normal')
                p.paragraph_format.left_indent = Inches(0.5)
                p.add_run(child.string.strip())
    
    # Process any other child elements recursively
    elif tag_name not in ['ul', 'ol', 'table', 'pre', 'blockquote']:
        # Safely process children by checking if they exist
        if hasattr(element, 'contents'):
            for child in element.contents:
                _process_element(child, document)


def _process_inline_elements(html_element, paragraph):
    """Process inline formatting within paragraphs."""
    for content in html_element.contents:
        if content.name:
            # This is a tag
            tag_name = content.name
            text = content.get_text()
            
            if tag_name == 'strong' or tag_name == 'b':
                run = paragraph.add_run(text)
                run.bold = True
            elif tag_name == 'em' or tag_name == 'i':
                run = paragraph.add_run(text)
                run.italic = True
            elif tag_name == 'code':
                run = paragraph.add_run(text)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            elif tag_name == 'a':
                run = paragraph.add_run(text)
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 255)
            else:
                # Unknown tag, just add the text
                paragraph.add_run(text)
        else:
            # This is just text
            if content.string and content.string.strip():
                paragraph.add_run(content.string)


def _process_list(list_element, document):
    """Process lists (ordered and unordered)."""
    is_ordered = list_element.name == 'ol'
    
    for i, li in enumerate(list_element.find_all('li', recursive=False)):
        p = document.add_paragraph(style='ListStyle')
        prefix = f"{i+1}. " if is_ordered else "• "
        p.add_run(prefix)
        
        # Process the content of the list item
        for content in li.contents:
            if hasattr(content, 'name') and content.name:
                if content.name in ['ul', 'ol']:
                    # Nested list
                    _process_list(content, document)
                else:
                    # Other tags
                    _process_inline_elements(content, p)
            else:
                # Text content
                if content.string and content.string.strip():
                    p.add_run(content.string)


def _process_table(table_element, document):
    """Process HTML tables."""
    # Get rows
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Count columns based on the first row
    header_cells = rows[0].find_all(['th', 'td'])
    col_count = len(header_cells)
    
    # Create the table
    table = document.add_table(rows=1, cols=col_count)
    table.style = 'CustomTableStyle'
    
    # Process header row
    header_row = table.rows[0]
    for i, cell in enumerate(header_cells):
        if i < col_count:  # Protection against malformed tables
            header_row.cells[i].text = cell.get_text().strip()
            header_row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in header_row.cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.name = 'Calibri'
    
    # Process data rows
    for row_idx, row in enumerate(rows[1:], 1):
        data_cells = row.find_all('td')
        if not data_cells:
            continue
            
        # Add a new row
        table_row = table.add_row()
        
        # Fill the cells
        for i, cell in enumerate(data_cells):
            if i < col_count:  # Protection against malformed tables
                table_row.cells[i].text = cell.get_text().strip()
                table_row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in table_row.cells[i].paragraphs[0].runs:
                    run.font.name = 'Calibri'
    
    # Format table
    for row in table.rows:
        row.height = Pt(24)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Apply consistent formatting to paragraphs in cells
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
